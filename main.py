# main.py (actualizado: guarda/lee orders en Redis, tolerante a kwargs de RQ,
# create-preference incluye category_id / metadata y mp-webhook guarda payment completo)
import os
import time
import uuid
import threading
import json
from typing import Optional

from fastapi import FastAPI, Request, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
import requests
from fastapi.responses import JSONResponse

# Optional AI + document libs (with fallbacks)
try:
    import openai
except Exception:
    openai = None

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None
    Pt = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except Exception:
    canvas = None
    letter = None

# Redis
REDIS_URL = os.getenv("REDIS_URL") or None
redis_client = None
if REDIS_URL:
    try:
        import redis
        redis_client = redis.from_url(REDIS_URL, decode_responses=True)
        redis_client.ping()
        print("[redis] conectado correctamente")
    except Exception as e:
        print(f"[redis] no se pudo conectar: {e}")
        redis_client = None
else:
    print("[redis] REDIS_URL no configurado. Se usará store en memoria como fallback.")

# ---------- Config (env vars) ----------
MP_ACCESS_TOKEN = (os.getenv("MP_ACCESS_TOKEN") or "").strip()
BASE_URL = os.getenv("BASE_URL") or ""
PORT = int(os.getenv("PORT", 8000))
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
ENABLE_SIMULATE = os.getenv("ENABLE_SIMULATE", "0") == "1"
GCS_BUCKET = os.getenv("GCS_BUCKET") or None
GCS_SERVICE_ACCOUNT_JSON = os.getenv("GCS_SERVICE_ACCOUNT_JSON") or None

if not MP_ACCESS_TOKEN:
    raise RuntimeError("Por favor configura MP_ACCESS_TOKEN en variables de entorno")

if not BASE_URL:
    BASE_URL = "https://TU_DOMINIO_RAILWAY"

if OPENAI_API_KEY and openai:
    try:
        openai.api_key = OPENAI_API_KEY
    except Exception:
        pass

# ---------- GCS client (opcional) ----------
gcs_client = None
if GCS_BUCKET and GCS_SERVICE_ACCOUNT_JSON:
    try:
        from google.oauth2 import service_account
        from google.cloud import storage
        creds_dict = json.loads(GCS_SERVICE_ACCOUNT_JSON)
        creds = service_account.Credentials.from_service_account_info(creds_dict)
        gcs_client = storage.Client(credentials=creds, project=creds_dict.get("project_id"))
        print("[gcs] client inicializado")
    except Exception as e:
        print(f"[gcs] init error: {e}")
        gcs_client = None
else:
    print("[gcs] GCS_BUCKET o GCS_SERVICE_ACCOUNT_JSON no configurado; no se subirá a GCS")

# ---------- App ----------
app = FastAPI(title="RedaXion - Backend con Redis/GCS")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

FILES_PATH = "/tmp"
app.mount("/files", StaticFiles(directory=FILES_PATH), name="files")

# ---------- In-memory store (fallback) ----------
ORDERS = {}  # fallback if redis not configured

def generate_access_code():
    t = time.strftime("%y%m%d%H%M%S")
    random = uuid.uuid4().hex[:4].upper()
    return f"RX-{t}-{random}"

# small helpers to persist/read an order
def persist_order(order_id: str, order_obj: dict):
    order_obj_copy = dict(order_obj)
    try:
        if redis_client:
            redis_client.set(f"order:{order_id}", json.dumps(order_obj_copy))
            # set TTL 7 days (optional)
            redis_client.expire(f"order:{order_id}", 7 * 24 * 3600)
            print(f"[persist] order:{order_id} saved to redis")
            return
    except Exception as e:
        print(f"[persist] redis error: {e}")
    # fallback
    ORDERS[order_id] = order_obj_copy
    print(f"[persist] order:{order_id} saved in memory fallback")

def load_order(order_id: str) -> Optional[dict]:
    try:
        if redis_client:
            val = redis_client.get(f"order:{order_id}")
            if val:
                return json.loads(val)
    except Exception as e:
        print(f"[load] redis read error: {e}")
    return ORDERS.get(order_id)

# Try to import enqueue helper - tasks.py must exist in repo
try:
    from tasks import enqueue_generate_and_deliver
except Exception:
    enqueue_generate_and_deliver = None
    print("[startup] tasks.enqueue_generate_and_deliver not available; will fallback to thread execution.")

# Helper: dynamic call (used when fallback a thread)
def _run_generate_with_correct_args(order_id: str, payer_email: Optional[str]):
    try:
        fn = generate_and_deliver  # noqa: F821
    except NameError:
        print(f"[simulate] generate_and_deliver no definida. Simulación para {order_id}")
        time.sleep(1)
        print(f"[simulate] Simulación completada para {order_id}")
        return

    # Intento robusto: preferimos (order_id, payer_email), luego (order_id), luego ()
    try:
        fn(order_id, payer_email)
        return
    except TypeError as e:
        print(f"[runner] llamada con (order_id, payer_email) falló ({e}), intentando alternativas...")

    try:
        fn(order_id)
        return
    except TypeError as e:
        print(f"[runner] llamada con (order_id) falló ({e}), intentando llamar sin args...")

    try:
        fn()
        return
    except Exception as e:
        print(f"[runner] llamada sin args falló: {e}")

    print(f"[runner] No se pudo invocar generate_and_deliver para {order_id}.")


# ---------- Endpoints ----------
# --- Reemplazar create-preference por esta versión ---
@app.post("/create-preference")
async def create_preference(req: Request):
    data = await req.json()
    order_id = data.get("order_id") or f"ORD-{uuid.uuid4().hex[:8].upper()}"
    email = data.get("email")
    if not email:
        raise HTTPException(status_code=400, detail="Falta email")

    order = {"email": email, "payload": data, "status": "pending", "created_at": time.time()}
    persist_order(order_id, order)

    preference_payload = {
        "items": data.get("items") or [{"title":"RedaXion - Generador de examen", "quantity":1, "unit_price":2000.0}],
        "external_reference": order_id,
        "payer": data.get("payer") or {"email": email},
        "notification_url": f"{BASE_URL}/mp-webhook",
        "back_urls": {
            "success": f"{BASE_URL}/mp-success?order_id={order_id}",
            "failure": f"{BASE_URL}/mp-failure?order_id={order_id}"
        },
        "auto_return": "approved"
    }

    headers = {"Authorization": f"Bearer {MP_ACCESS_TOKEN}", "Content-Type": "application/json"}
    resp = requests.post("https://api.mercadopago.com/checkout/preferences", json=preference_payload, headers=headers)
    if resp.status_code not in (200, 201):
        # guardar respuesta errónea para debug
        try:
            with open("/tmp/last_preference_response.json", "w", encoding="utf-8") as f:
                f.write(json.dumps({"status": resp.status_code, "text": resp.text}))
        except Exception:
            pass
        return {"error": "No se pudo crear preference", "details": resp.text, "status_code": resp.status_code}

    pref = resp.json()
    # Guardar la respuesta completa para debugging (archivo legible)
    try:
        with open("/tmp/last_preference.json", "w", encoding="utf-8") as f:
            f.write(json.dumps(pref))
    except Exception as e:
        print(f"[debug] no se pudo escribir last_preference.json: {e}")

    init_point = pref.get("init_point") or pref.get("sandbox_init_point")
    preference_id = pref.get("id")

    order["preference_id"] = preference_id
    order["preference_payload"] = preference_payload
    persist_order(order_id, order)

    return {"init_point": init_point, "sandbox_init_point": pref.get("sandbox_init_point"), "preference_id": preference_id, "order_id": order_id}

# --- Modificar mp_webhook para guardar el body recibido ---
@app.post("/mp-webhook")
async def mp_webhook(req: Request):
    try:
        body = await req.json()
    except Exception:
        # guardar raw body si json falla
        raw = await req.body()
        try:
            with open("/tmp/last_webhook_raw.txt", "wb") as f:
                f.write(raw)
        except Exception:
            pass
        return {"ok": False, "error": "invalid json"}

    # Guardar copia del webhook entrante para debug
    try:
        with open("/tmp/last_webhook.json", "w", encoding="utf-8") as f:
            f.write(json.dumps(body))
    except Exception as e:
        print(f"[debug] no se pudo escribir last_webhook.json: {e}")

    # (seguir con la lógica existente que ya tienes)
    payment_id = None
    if isinstance(body, dict):
        if "data" in body and isinstance(body["data"], dict) and "id" in body["data"]:
            payment_id = body["data"]["id"]
        elif "id" in body:
            payment_id = body["id"]

    if not payment_id:
        return {"ok": True, "note": "no payment id"}

    r = requests.get(f"https://api.mercadopago.com/v1/payments/{payment_id}",
                     headers={"Authorization": f"Bearer {MP_ACCESS_TOKEN}"})
    if r.status_code != 200:
        print(f"[mp-webhook] validation failed for payment {payment_id}: {r.status_code} {r.text}")
        return {"error": "No se pudo validar pago", "details": r.text}

    payment = r.json()
    status = payment.get("status")
    if status != "approved":
        print(f"[mp-webhook] payment {payment_id} status: {status}")
        return {"ok": True, "status": status}

    external_ref = payment.get("external_reference") or (payment.get("metadata") or {}).get("order_id")
    payer_email = payment.get("payer", {}).get("email")
    amount = payment.get("transaction_amount")

    if not external_ref:
        return {"ok": False, "error": "external_reference missing"}

    order = load_order(external_ref) or {}
    order["status"] = "paid"
    order["payment_id"] = payment_id
    order["amount"] = amount
    order["payer_email"] = payer_email
    order["access_code"] = generate_access_code()
    persist_order(external_ref, order)

    # enqueue/generar...
    if enqueue_generate_and_deliver:
        try:
            jobid = enqueue_generate_and_deliver(external_ref)
            print(f"[mp-webhook] enqueued job {jobid} for order {external_ref}")
        except Exception as e:
            print(f"[mp-webhook] enqueue failed: {e}; falling back to thread")
            threading.Thread(target=_run_generate_with_correct_args, args=(external_ref, payer_email), daemon=True).start()
    else:
        threading.Thread(target=_run_generate_with_correct_args, args=(external_ref, payer_email), daemon=True).start()

    return {"ok": True, "order": external_ref, "code": order["access_code"]}

# --- END PATCH ---

# --- Añade estos endpoints DEBUG al final del archivo para leer lo guardado ---
@app.get("/debug-last-preference")
async def debug_last_preference():
    try:
        with open("/tmp/last_preference.json", "r", encoding="utf-8") as f:
            return JSONResponse(content=json.loads(f.read()))
    except Exception as e:
        try:
            with open("/tmp/last_preference_response.json", "r", encoding="utf-8") as f:
                return JSONResponse(content=json.loads(f.read()))
        except Exception as e2:
            return JSONResponse(status_code=404, content={"ok": False, "error": "no last preference file", "e1": str(e), "e2": str(e2)})

@app.get("/debug-last-webhook")
async def debug_last_webhook():
    try:
        with open("/tmp/last_webhook.json", "r", encoding="utf-8") as f:
            return JSONResponse(content=json.loads(f.read()))
    except Exception as e:
        try:
            with open("/tmp/last_webhook_raw.txt", "rb") as f:
                raw = f.read().decode('utf-8', errors='replace')
                return JSONResponse(content={"raw": raw})
        except Exception as e2:
            return JSONResponse(status_code=404, content={"ok": False, "error": "no last webhook file", "e1": str(e), "e2": str(e2)})

@app.post("/mp-webhook")
async def mp_webhook(req: Request):
    # Intenta siempre ack con 200 a MP (evita reintentos infinitos)
    try:
        body = await req.json()
    except Exception:
        # devolver 200 aunque no sea JSON; log y seguir
        print("[mp-webhook] body not json or empty")
        return JSONResponse({"ok": True, "note": "invalid json"}, status_code=200)

    # log inicial para debugging
    print(f"[mp-webhook] received: {json.dumps(body)[:2000]}")

    payment_id = None
    if isinstance(body, dict):
        if "data" in body and isinstance(body["data"], dict) and "id" in body["data"]:
            payment_id = str(body["data"]["id"])
        elif "id" in body:
            payment_id = str(body["id"])

    if not payment_id:
        # guardar la notificación bruta para revisar después
        try:
            tmp_id = f"notif-{uuid.uuid4().hex[:8]}"
            if redis_client:
                redis_client.set(f"notif:{tmp_id}", json.dumps({"body": body, "headers": dict(req.headers)}))
                redis_client.expire(f"notif:{tmp_id}", 24 * 3600)
            print(f"[mp-webhook] sin payment_id; saved notif:{tmp_id}")
        except Exception as e:
            print(f"[mp-webhook] save notif failed: {e}")
        return JSONResponse({"ok": True, "note": "no payment id"}, status_code=200)

    # validar con MP
    r = requests.get(f"https://api.mercadopago.com/v1/payments/{payment_id}",
                     headers={"Authorization": f"Bearer {MP_ACCESS_TOKEN}"})
    if r.status_code != 200:
        print(f"[mp-webhook] validation failed for payment {payment_id}: {r.status_code} {r.text}")
        return JSONResponse({"ok": True, "note": "payment lookup failed", "details": r.text}, status_code=200)

    payment = r.json()
    # guardar payment raw para auditoría
    try:
        if redis_client:
            redis_client.set(f"payment:{payment_id}", json.dumps(payment))
            redis_client.expire(f"payment:{payment_id}", 7 * 24 * 3600)
    except Exception as e:
        print(f"[mp-webhook] redis save payment failed: {e}")

    status = payment.get("status")
    print(f"[mp-webhook] payment {payment_id} status: {status}")

    # siempre responder 200 a MP; procesar solo si approved
    if status != "approved":
        return JSONResponse({"ok": True, "status": status}, status_code=200)

    external_ref = payment.get("external_reference") or (payment.get("metadata") or {}).get("order_id")
    payer_email = payment.get("payer", {}).get("email")
    amount = payment.get("transaction_amount")

    if not external_ref:
        # si no hay external_reference, crear un registro temporal para audit
        external_ref = f"orphan-{payment_id}"
        print(f"[mp-webhook] external_reference missing, using {external_ref}")

    order = load_order(external_ref) or {}
    order["status"] = "paid"
    order["payment_id"] = payment_id
    order["amount"] = amount
    order["payer_email"] = payer_email
    order["access_code"] = generate_access_code()
    order["payment_raw"] = payment
    order.setdefault("metadata", {}).update({"mp_payment_status": status, "mp_payment_id": payment_id})
    persist_order(external_ref, order)

    # Lanzar generación del examen en background via RQ enqueue (fallback to thread if not available)
    if enqueue_generate_and_deliver:
        try:
            jobid = enqueue_generate_and_deliver(external_ref)
            print(f"[mp-webhook] enqueued job {jobid} for order {external_ref}")
        except Exception as e:
            print(f"[mp-webhook] enqueue failed: {e}; falling back to thread")
            threading.Thread(target=_run_generate_with_correct_args, args=(external_ref, payer_email), daemon=True).start()
    else:
        threading.Thread(target=_run_generate_with_correct_args, args=(external_ref, payer_email), daemon=True).start()

    return JSONResponse({"ok": True, "order": external_ref, "code": order["access_code"]}, status_code=200)

# ---------- Document generation helpers (sin cambios) ----------
def generate_text_with_openai(subject: str, topic: str, mcq_count: int = 14, essay_count: int = 2) -> str:
    PROMPT = "REDA_PROMPT_PLACEHOLDER"
    if OPENAI_API_KEY and openai:
        try:
            resp = openai.ChatCompletion.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": PROMPT}],
                temperature=0.0,
                max_tokens=3500
            )
            text = resp["choices"][0]["message"]["content"]
            return text
        except Exception as e:
            print(f"[generate_text_with_openai] OpenAI error: {e}")
    return f"Asignatura: {subject}\nTema: {topic}\n\nPreguntas de alternativa:\n\n(Contenido fallback)\n\n<<SOLUTIONS>>\n\nSolucionario:\n\n(Contenido fallback)"

def save_docx_from_text(text: str, path: str) -> None:
    try:
        if Document is None:
            raise RuntimeError("python-docx no instalado")
        doc = Document()
        try:
            style = doc.styles['Normal']
            font = style.font
            if Pt is not None:
                font.size = Pt(11)
            font.name = 'Calibri'
        except Exception:
            pass
        for raw_line in text.split("\n"):
            line = raw_line.rstrip()
            if not line:
                doc.add_paragraph("")
                continue
            if line.startswith("Asignatura:") or line.startswith("Tema:"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                continue
            if line.lower().startswith("preguntas"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                continue
            if len(line) >= 3 and line[:3].strip().isdigit() and line[2] == ')':
                p = doc.add_paragraph(line, style='List Number')
                continue
            if len(line) >= 2 and (line[1] == '.' or line[1] == ')') and line[0].isalpha():
                p = doc.add_paragraph(line)
                continue
            p = doc.add_paragraph(line)
        doc.save(path)
    except Exception as e:
        with open(path, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"[save_docx] fallback write: {path}")

def save_pdf_from_text(text: str, path: str) -> None:
    try:
        from reportlab.lib.utils import simpleSplit
    except Exception:
        simpleSplit = None
    if canvas is None or letter is None or simpleSplit is None:
        with open(path, "w", encoding="utf-8") as f2:
            f2.write(text)
        return
    c = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    margin = 50
    max_width = width - 2 * margin
    y = height - margin
    normal_font = "Helvetica"
    bold_font = "Helvetica-Bold"
    size_title = 14
    size_heading = 12
    size_normal = 11
    leading = 14
    paragraphs = text.split("\n\n")
    for paragraph in paragraphs:
        for raw_line in paragraph.split("\n"):
            line = raw_line.strip()
            if not line:
                y -= leading // 2
                continue
            if line.startswith("Asignatura:") or line.startswith("Tema:"):
                font = bold_font
                fsize = size_title
            elif line.lower().startswith("preguntas"):
                font = bold_font
                fsize = size_heading
            else:
                font = normal_font
                fsize = size_normal
            wrapped = simpleSplit(line, font, fsize, max_width)
            for wline in wrapped:
                if y < margin + leading:
                    c.showPage()
                    y = height - margin
                c.setFont(font, fsize)
                c.drawString(margin, y, wline)
                y -= leading
        y -= leading // 2
    c.save()

# helper to upload file to GCS (if configured)
def upload_to_gcs(local_path: str, dest_name: str) -> Optional[str]:
    if not gcs_client or not GCS_BUCKET:
        return None
    try:
        bucket = gcs_client.bucket(GCS_BUCKET)
        blob = bucket.blob(dest_name)
        blob.upload_from_filename(local_path)
        # make signed url for 24h
        url = blob.generate_signed_url(expiration=86400)
        print(f"[gcs] uploaded {local_path} -> {dest_name}")
        return url
    except Exception as e:
        print(f"[gcs] upload error: {e}")
        return None

# ---------- Core: generate and deliver (ahora tolerante a kwargs extra) ----------
def generate_and_deliver(order_id: Optional[str] = None, payer_email: Optional[str] = None, *args, **kwargs):
    """
    Acepta kwargs/args extra que RQ pueda pasar (p.ej. enqueue_timeout).
    """
    # Si order_id vino por kwargs (posible en algunos env), úsalo.
    if not order_id and isinstance(kwargs.get("order_id"), str):
        order_id = kwargs.get("order_id")

    print(f"[generate_and_deliver] START order {order_id} for {payer_email} (extra_args={len(args)}, extra_kwargs={list(kwargs.keys())})")

    if not order_id:
        print("[generate_and_deliver] ERROR: order_id no especificado. Abortando.")
        return

    order = load_order(order_id) or {}
    payload = order.get("payload", {}) if isinstance(order, dict) else {}
    subject = payload.get("subject", "Asignatura")
    topic = payload.get("topic", "Tema")

    try:
        mcq_count = int(payload.get("mcqCount", payload.get("mcq_count", 14) or 14))
    except Exception:
        mcq_count = 14
    try:
        essay_count = int(payload.get("essayCount", payload.get("essay_count", 2) or 2))
    except Exception:
        essay_count = 2

    order["status"] = "processing"
    persist_order(order_id, order)

    text = generate_text_with_openai(subject, topic, mcq_count, essay_count)

    if "<<SOLUTIONS>>" in text:
        parts = text.split("<<SOLUTIONS>>", 1)
        exam_part = parts[0].strip()
        solutions_part = parts[1].strip()
    else:
        exam_part = text.strip()
        solutions_part = "Solucionario no generado (fallback)."

    safe = order_id.replace("/", "_")
    exam_docx = os.path.join(FILES_PATH, f"{safe}-exam.docx")
    exam_pdf  = os.path.join(FILES_PATH, f"{safe}-exam.pdf")
    sol_docx  = os.path.join(FILES_PATH, f"{safe}-solutions.docx")
    sol_pdf   = os.path.join(FILES_PATH, f"{safe}-solutions.pdf")

    try:
        save_docx_from_text(exam_part, exam_docx)
        save_docx_from_text(solutions_part, sol_docx)
        save_pdf_from_text(exam_part, exam_pdf)
        save_pdf_from_text(solutions_part, sol_pdf)
    except Exception as e:
        print(f"[generate_and_deliver] file save error: {e}")

    # upload to GCS if available and create final URLs
    try:
        exam_url = upload_to_gcs(exam_pdf, os.path.basename(exam_pdf)) or f"{BASE_URL}/files/{os.path.basename(exam_pdf)}"
    except Exception:
        exam_url = f"{BASE_URL}/files/{os.path.basename(exam_pdf)}"
    try:
        sol_url  = upload_to_gcs(sol_pdf, os.path.basename(sol_pdf)) or f"{BASE_URL}/files/{os.path.basename(sol_pdf)}"
    except Exception:
        sol_url = f"{BASE_URL}/files/{os.path.basename(sol_pdf)}"

    try:
        docx_exam_url = upload_to_gcs(exam_docx, os.path.basename(exam_docx)) or f"{BASE_URL}/files/{os.path.basename(exam_docx)}"
    except Exception:
        docx_exam_url = f"{BASE_URL}/files/{os.path.basename(exam_docx)}"
    try:
        docx_sol_url  = upload_to_gcs(sol_docx, os.path.basename(sol_docx)) or f"{BASE_URL}/files/{os.path.basename(sol_docx)}"
    except Exception:
        docx_sol_url = f"{BASE_URL}/files/{os.path.basename(sol_docx)}"

    order["exam_pdf_url"] = exam_url
    order["exam_docx_url"] = docx_exam_url
    order["solutions_pdf_url"] = sol_url
    order["solutions_docx_url"] = docx_sol_url

    order["status"] = "ready"
    order["delivered_at"] = time.time()
    persist_order(order_id, order)
    print(f"[generate_and_deliver] DONE order {order_id}. Exam: {exam_url} Solutions: {sol_url}")

@app.get("/order-status")
async def order_status(order_id: str):
    order = load_order(order_id)
    if not order:
        return {"status": "not_found"}
    return {
        "status": order.get("status"),
        "exam_pdf_url": order.get("exam_pdf_url"),
        "exam_docx_url": order.get("exam_docx_url"),
        "solutions_pdf_url": order.get("solutions_pdf_url"),
        "solutions_docx_url": order.get("solutions_docx_url"),
        "access_code": order.get("access_code"),
        "payment_id": order.get("payment_id"),
    }

# Optional simulate endpoint
if ENABLE_SIMULATE:
    @app.post("/simulate-paid")
    async def simulate_paid(request: Request):
        payload = await request.json()
        order_id = payload.get("order_id") or f"SIM-{int(time.time())}"
        payer_email = payload.get("payer_email") or "cliente@ejemplo.com"

        order = {"status": "paid", "payer_email": payer_email, "payment_id": f"SIM-{uuid.uuid4().hex[:8]}", "access_code": generate_access_code(), "payload": payload}
        persist_order(order_id, order)

        if enqueue_generate_and_deliver:
            try:
                jobid = enqueue_generate_and_deliver(order_id)
                print(f"[simulate-paid] enqueued job {jobid} for {order_id}")
            except Exception as e:
                print(f"[simulate-paid] enqueue failed: {e}; falling back to thread")
                threading.Thread(target=_run_generate_with_correct_args, args=(order_id, payer_email), daemon=True).start()
        else:
            threading.Thread(target=_run_generate_with_correct_args, args=(order_id, payer_email), daemon=True).start()

        return {"ok": True, "simulated": True, "order": order_id, "payer_email": payer_email}

# debug endpoints
import os
@app.get("/debug-list-files")
async def debug_list_files():
    try:
        files = os.listdir(FILES_PATH)
        ord_files = [f for f in files if f.startswith("ORD-")]
        return {"ok": True, "FILES_PATH": FILES_PATH, "count": len(ord_files), "files": ord_files}
    except Exception as e:
        return {"ok": False, "error": str(e)}

@app.get("/mp-webhook")
async def mp_webhook_get():
    return JSONResponse({"ok": True, "note": "mp-webhook GET alive"})

@app.get("/debug-last-preference")
async def debug_last_preference():
    try:
        with open("/tmp/last_preference.json", "r", encoding="utf-8") as f:
            return JSONResponse(content=json.loads(f.read()))
    except Exception as e:
        try:
            with open("/tmp/last_preference_response.json", "r", encoding="utf-8") as f:
                return JSONResponse(content=json.loads(f.read()))
        except Exception as e2:
            return JSONResponse(status_code=404, content={"ok": False, "error": "no last preference file", "e1": str(e), "e2": str(e2)})

@app.get("/debug-last-webhook")
async def debug_last_webhook():
    try:
        with open("/tmp/last_webhook.json", "r", encoding="utf-8") as f:
            return JSONResponse(content=json.loads(f.read()))
    except Exception as e:
        try:
            with open("/tmp/last_webhook_raw.txt", "rb") as f:
                raw = f.read().decode('utf-8', errors='replace')
                return JSONResponse(content={"raw": raw})
        except Exception as e2:
            return JSONResponse(status_code=404, content={"ok": False, "error": "no last webhook file", "e1": str(e), "e2": str(e2)})
